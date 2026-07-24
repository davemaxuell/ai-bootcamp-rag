from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "submission" / "BUFS_Admissions_RAG_결과보고서_20232829.docx"

AUTHOR = "맥슈웰 데이브"
STUDENT_ID = "20232829"
REPORT_DATE = "2026년 7월 24일"
GITHUB_URL = "https://github.com/davemaxuell/ai-bootcamp-rag"
DEPLOY_URL = "https://ai-bootcamp-rag-7ahkiva6fmc7qzkiifmz6a.streamlit.app/"

CALIBRI = "Calibri"
KOREAN_FONT = "Malgun Gothic"
NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66717D"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D8DEE5"
PALE_BLUE = "EEF5FB"
GOLD = "A66F00"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    size=None,
    color=None,
    bold=None,
    italic=None,
    ascii_font=CALIBRI,
    east_asia_font=KOREAN_FONT,
):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=None, bold=None):
    style.font.name = CALIBRI
    style._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, color="222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    set_style_font(title, 30, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 15, color=DARK_BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, 11, color="222222")
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = doc.styles["Caption"]
    set_style_font(caption, 9, color=MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def configure_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "BUFS MULTILINGUAL ADMISSIONS RAG  |  FINAL PROJECT REPORT"
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    prefix = paragraph.add_run(f"{AUTHOR} · {STUDENT_ID}   |   ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(paragraph)


def paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color, size=18, space=10):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pbdr.append(left)


def add_lead_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph_shading(paragraph, PALE_BLUE)
    paragraph_left_border(paragraph, BLUE)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color="222222")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=MID_GRAY, size=6):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {PAGE_WIDTH_DXA}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblpr = tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def format_cell_text(cell, *, bold=False, color="222222", size=9.5, align=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        shade_cell(header_cells[idx], LIGHT_GRAY)
        format_cell_text(
            header_cells[idx],
            bold=True,
            color=NAVY,
            size=9.5,
            align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT),
        )
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = str(value)
            format_cell_text(
                cell,
                size=9.3,
                align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT),
            )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.add_run(text)
    return paragraph


def add_number(doc, title, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, bold=True, color=DARK_BLUE)
    paragraph.add_run(text)
    return paragraph


def add_metadata_line(doc, label, value, *, align=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=MUTED, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5, color=NAVY)


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_report():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "BUFS 다국어 입학 도우미 프로젝트 결과보고서"
    doc.core_properties.subject = "LangChain 및 RAG 기반 다국어 실용 서비스"
    doc.core_properties.author = f"{AUTHOR} ({STUDENT_ID})"
    doc.core_properties.keywords = "LangChain, RAG, Chroma, Streamlit, multilingual"
    doc.core_properties.comments = "부산외국어대학교 AI 부트캠프 최종 프로젝트"

    configure_styles(doc)
    for section in doc.sections:
        set_page_geometry(section)
        configure_header_footer(section)

    # Page 1 — Editorial cover.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("MULTILINGUAL AI SERVICE · FINAL PROJECT")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BUFS 다국어 입학 도우미")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("LangChain & RAG 기반 외국인 지원자 맞춤형 입학 안내 서비스")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(54)
    line_run = line.add_run("Official documents → semantic retrieval → grounded multilingual answers")
    set_run_font(line_run, size=10, color=MUTED, italic=True)

    add_metadata_line(doc, "작성자", AUTHOR)
    add_metadata_line(doc, "학번", STUDENT_ID)
    add_metadata_line(doc, "제출일", REPORT_DATE)
    add_metadata_line(doc, "GitHub", GITHUB_URL)
    add_metadata_line(doc, "Web", DEPLOY_URL)

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(42)
    run = closing.add_run("부산외국어대학교 · 첨단산업 인재양성 AI 부트캠프")
    set_run_font(run, size=10, color=MUTED)

    add_page_break(doc)

    # Page 2 — Executive summary and requirements.
    doc.add_heading("1. 프로젝트 개요", level=1)
    add_lead_callout(
        doc,
        "핵심 성과",
        "한국어·영어 공식 입학 문서 6개를 64개 검색 단위로 구축하고, 질문 언어를 유지한 근거 기반 Q&A와 개인화 제출서류 체크리스트를 하나의 Streamlit 서비스로 구현했다.",
    )

    paragraph = doc.add_paragraph()
    paragraph.add_run(
        "외국인 지원자는 모집요강, 온라인 접수 안내, 제출서류 체크리스트에 흩어진 정보를 "
        "여러 언어로 이해해야 한다. 일반 번역기는 문장을 바꾸는 데 그치며, 지원 과정과 "
        "신입·편입 여부에 따라 필요한 정보를 찾아 조합하고 공식 근거를 제시하지 못한다. "
        "본 프로젝트는 RAG로 이 문제를 해결한다."
    )

    doc.add_heading("1.1 목표", level=2)
    add_bullet(doc, "사용자가 한국어, 영어, 중국어, 베트남어 등 원하는 언어로 질문할 수 있게 한다.")
    add_bullet(doc, "답변을 공식 BUFS 입학 문서의 파일명과 페이지에 연결하여 검증 가능하게 한다.")
    add_bullet(doc, "지원 과정과 지원 유형을 반영한 제출서류 체크리스트를 자동 생성한다.")
    add_bullet(doc, "근거가 없으면 추측하지 않고 입학처 확인을 안내하는 안전장치를 제공한다.")

    doc.add_heading("1.2 구현 범위 및 성과", level=2)
    add_table(
        doc,
        ["구분", "구현 결과"],
        [
            ("데이터", "BUFS 공식 입학 PDF 6개 · 한국어/영어"),
            ("검색 인덱스", "Chroma VectorDB · 64개 임베딩 청크"),
            ("사용자 기능", "다국어 Q&A · 맞춤형 제출서류 체크리스트"),
            ("검증", "단위 테스트 12/12 · 다국어 라이브 평가 4/4"),
            ("웹 서비스", "Streamlit Chat/Checklist 2개 탭"),
        ],
        [2300, 7060],
    )

    doc.add_heading("1.3 과제 요구사항 대응", level=2)
    add_table(
        doc,
        ["요구사항", "적용 내용", "상태"],
        [
            ("Claude Code", "요구 분석, 모듈 설계, 테스트·오류 수정·배포 워크플로에 활용", "충족"),
            ("LangChain", "Document, OpenAIEmbeddings, ChatOpenAI, Chroma 연동", "충족"),
            ("RAG", "다국어 문서 추출→임베딩→검색→근거 기반 생성", "충족"),
            ("고차원 기능", "출처 Q&A, 답변 거절, 지원자별 체크리스트", "충족"),
            ("웹 배포", "GitHub 기반 Streamlit Community Cloud 배포", "완료"),
        ],
        [2100, 5760, 1500],
        [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )

    # Page 3 — Architecture.
    doc.add_heading("2. 시스템 설계", level=1)
    doc.add_paragraph(
        "서비스는 사전 구축 단계와 사용자 요청 단계로 분리된다. 사전 구축에서는 PDF를 "
        "추출·임베딩하여 영속 Chroma 인덱스를 만들고, 요청 단계에서는 질문마다 관련 문맥을 "
        "검색해 LLM에 제한된 근거로 제공한다."
    )

    doc.add_heading("2.1 데이터 흐름", level=2)
    add_number(doc, "PDF 추출", "`pdfplumber`를 우선 사용하고 짧은 페이지는 `pypdf`로 보완한다.")
    add_number(doc, "청크 분할", "페이지 정보를 유지하면서 최대 2,000자 단위로 문서를 분할한다.")
    add_number(doc, "메타데이터", "source, page, lang, level을 각 청크에 저장한다.")
    add_number(doc, "임베딩", "`text-embedding-3-small`로 의미 벡터를 생성한다.")
    add_number(doc, "검색", "질문과 유사한 상위 문서를 찾고 지원 과정 필터를 적용한다.")
    add_number(doc, "생성", "`gpt-4o-mini`가 검색 문맥만 사용하여 질문 언어로 답변한다.")

    doc.add_heading("2.2 주요 컴포넌트", level=2)
    add_table(
        doc,
        ["모듈", "역할", "핵심 설계"],
        [
            ("extract.py", "PDF 텍스트 추출·분할", "페이지와 파일 출처 보존, Windows 중복 glob 방지"),
            ("ingest.py", "VectorDB 구축", "안정적 SHA-256 ID로 재인덱싱 시 중복 방지"),
            ("retriever.py", "유사도 검색", "학부/대학원 + 공통 문서 메타데이터 필터"),
            ("answerer.py", "근거 기반 답변", "동일 언어 응답, 인라인 출처, 빈 문맥 즉시 거절"),
            ("checklist.py", "맞춤 체크리스트", "과정·전공·신입/편입·언어를 프롬프트에 반영"),
            ("app.py", "웹 인터페이스", "질문 탭과 체크리스트 탭, 오류 메시지와 출처 표시"),
        ],
        [1700, 2600, 5060],
    )

    doc.add_heading("2.3 다국어 RAG 전략", level=2)
    doc.add_paragraph(
        "원문 데이터는 한국어와 영어지만, 다국어 의미 공간을 지원하는 임베딩 모델을 사용해 "
        "중국어·베트남어 질문도 관련 한국어/영어 문서로 연결한다. 답변 시에는 시스템 지침으로 "
        "사용자의 질문 언어를 유지한다. 즉, 번역 자체가 목적이 아니라 문서 검색, 맥락 결합, "
        "개인화된 정보 구조화가 핵심이다."
    )

    # Page 4 — Features and safeguards.
    doc.add_heading("3. 핵심 기능", level=1)

    doc.add_heading("3.1 공식 문서 기반 다국어 Q&A", level=2)
    doc.add_paragraph(
        "사용자가 어느 언어로 질문하더라도 관련 모집요강과 안내서의 상위 문맥을 검색한다. "
        "LLM은 제공된 문맥 밖의 날짜, 금액, 자격요건을 만들지 않도록 제한되며, 답변과 별도로 "
        "검색에 사용된 파일명 및 페이지 목록을 UI에 표시한다."
    )
    add_bullet(doc, "예시 질문: “What is the application deadline?”")
    add_bullet(doc, "예시 질문: “本科申请需要什么材料?”")
    add_bullet(doc, "예시 질문: “Cần những giấy tờ gì để nộp hồ sơ?”")

    doc.add_heading("3.2 개인화 제출서류 체크리스트", level=2)
    doc.add_paragraph(
        "사용자는 학부/대학원, 전공, 신입/편입, 출력 언어를 지정한다. 서비스는 일반 체크리스트와 "
        "해당 과정 모집요강을 함께 검색해 번호가 있는 서류 목록을 생성하고, 국적·학력·어학점수에 "
        "따라 달라지는 조건을 분리해서 표시한다."
    )

    doc.add_heading("3.3 신뢰성과 안전장치", level=2)
    add_table(
        doc,
        ["위험", "대응 방식"],
        [
            ("근거 없는 답변", "검색 문서가 없으면 LLM을 호출하지 않고 입학처 확인 문구 반환"),
            ("잘못된 과정 혼합", "학부/대학원 필터와 공통 문서를 함께 검색"),
            ("출처 불명확", "파일명·페이지를 문맥과 UI에 모두 표시"),
            ("인덱스 중복", "문서 내용 기반 해시 ID를 사용한 멱등적 ingest"),
            ("비밀키 노출", ".env/API_key.txt를 gitignore 처리하고 배포 Secrets 사용"),
        ],
        [2500, 6860],
    )

    doc.add_heading("3.4 웹 사용자 경험", level=2)
    add_bullet(doc, "Ask a question: 자연어 질문 입력 → 공식 문서 검색 → 다국어 답변 및 출처 표시")
    add_bullet(doc, "My document checklist: 지원자 조건 입력 → 개인화된 번호 목록 생성")
    add_bullet(doc, "API 키 미설정 및 실행 오류를 사용자가 이해할 수 있는 메시지로 표시")
    add_bullet(doc, "Streamlit 테마 설정으로 정보 중심의 일관된 화면 구성")

    doc.add_heading("3.5 Claude Code 활용", level=2)
    doc.add_paragraph(
        "Claude Code 기반 AI 코딩 워크플로를 활용해 요구사항을 구현 단위로 분해하고, 모듈별 "
        "인터페이스와 테스트를 먼저 정의한 뒤 구현·검증을 반복했다. PDF 중복 처리와 같은 "
        "Windows 환경 문제를 테스트 과정에서 발견해 수정했고, 배포 전 비밀키·git 상태·웹 "
        "헬스 체크를 점검하는 방식으로 개발 품질을 관리했다."
    )

    # Page 5 — Verification.
    doc.add_heading("4. 테스트 및 평가", level=1)
    add_lead_callout(
        doc,
        "검증 결과",
        "단위 테스트 12개와 한국어·영어·중국어·베트남어 라이브 평가 4개를 모두 통과했으며, Streamlit 로컬 헬스 엔드포인트에서 HTTP 200을 확인했다.",
    )

    doc.add_heading("4.1 단위 테스트", level=2)
    add_table(
        doc,
        ["검증 영역", "주요 확인 항목", "결과"],
        [
            ("설정", "모델명, 언어/과정 파일명 분류", "3/3"),
            ("PDF 추출", "청크 크기, 짧은 문서, 실제 메타데이터", "3/3"),
            ("Answerer", "빈 문맥 무호출 거절, 출처 포맷", "2/2"),
            ("Checklist", "검색 결과 없음 시 안전한 거절", "1/1"),
            ("Ingest", "안정적이고 내용 민감한 문서 ID", "1/1"),
            ("Retriever", "빈 질문 처리, 과정+공통 문서 필터", "2/2"),
        ],
        [2000, 5860, 1500],
        [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )

    doc.add_heading("4.2 다국어 라이브 평가", level=2)
    add_table(
        doc,
        ["언어", "질문 예시", "판정"],
        [
            ("한국어", "무슨 서류를 제출해야 하나요?", "PASS"),
            ("영어", "What is the application deadline?", "PASS"),
            ("중국어", "本科申请需要什么材料?", "PASS"),
            ("베트남어", "Cần những giấy tờ gì để nộp hồ sơ?", "PASS"),
        ],
        [1500, 6360, 1500],
        [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )

    doc.add_heading("4.3 실행 검증", level=2)
    add_bullet(doc, "실제 공식 PDF 6개에서 총 64개 텍스트 청크 생성")
    add_bullet(doc, "Persistent Chroma 인덱스에 벡터 64개 저장 확인")
    add_bullet(doc, "영어·중국어·베트남어 질의에서 관련 한국어/영어 문서 검색 확인")
    add_bullet(doc, "Streamlit 로컬 서버 시작 및 `/_stcore/health` HTTP 200 확인")
    add_bullet(doc, "공개 Streamlit URL에서 실제 질문의 답변과 파일명·페이지 출처 표시 확인")
    add_bullet(doc, "커밋 대상 파일에서 OpenAI 비밀키 패턴이 없음을 확인")

    doc.add_heading("4.4 평가 한계", level=2)
    doc.add_paragraph(
        "현재 골드 평가 세트는 대표 언어 4개를 빠르게 확인하는 스모크 테스트 수준이다. "
        "향후에는 마감일·등록금·지원자격·서류 예외조건 등 약 15개 이상의 정답 기반 질문으로 "
        "확장하고, 검색 정확도와 최종 답변 정확도를 별도로 측정할 필요가 있다."
    )

    # Page 6 — Deployment, limitations, conclusion.
    doc.add_heading("5. 배포 및 운영", level=1)
    doc.add_heading("5.1 제출·배포 링크", level=2)
    add_table(
        doc,
        ["항목", "링크/파일"],
        [
            ("GitHub 저장소", GITHUB_URL),
            ("최종 웹 서비스", DEPLOY_URL),
            ("소스 코드 노트북", "submission/BUFS_Admissions_RAG_20232829.ipynb"),
            ("프로젝트 결과보고서", "submission/BUFS_Admissions_RAG_결과보고서_20232829.docx"),
        ],
        [2600, 6760],
    )

    doc.add_heading("5.2 배포 구성", level=2)
    add_number(doc, "Repository", "소스 코드와 persistent `chroma_db/`를 GitHub `main` 브랜치에 저장한다.")
    add_number(doc, "Entry point", "Streamlit Community Cloud의 실행 파일을 `app.py`로 지정한다.")
    add_number(doc, "Secrets", "OpenAI 키는 저장소가 아닌 Streamlit Secrets의 `OPENAI_API_KEY`에 저장한다.")
    add_number(doc, "Operation", "GitHub에 새 커밋이 푸시되면 Streamlit 앱이 자동으로 재배포된다.")

    doc.add_heading("6. 한계 및 개선 방향", level=1)
    add_bullet(doc, "스크린샷 중심 페이지의 OCR은 v1 범위에서 제외되어 이미지 내부 텍스트를 검색하지 못한다.")
    add_bullet(doc, "PDF가 갱신되면 수동 ingest가 필요하므로 향후 자동 문서 동기화 파이프라인이 필요하다.")
    add_bullet(doc, "현재 모델 답변은 정보 안내 목적이며 중요한 일정과 자격요건은 입학처 재확인이 필요하다.")
    add_bullet(doc, "향후 검색 점수 임계값, reranker, 사용자 피드백을 추가하면 근거 적합도를 높일 수 있다.")

    doc.add_heading("7. 결론", level=1)
    doc.add_paragraph(
        "본 프로젝트는 단순한 1:1 번역기를 넘어, 서로 다른 언어로 작성된 공식 문서를 하나의 "
        "검색 가능한 지식베이스로 통합하고 지원자의 상황에 맞는 답변과 체크리스트를 제공한다. "
        "LangChain, Chroma VectorDB, 다국어 임베딩, LLM, Streamlit을 실제 서비스 흐름으로 "
        "연결했으며, 출처 표시와 답변 거절 정책으로 입학 정보 서비스에 필요한 신뢰성을 강화했다."
    )

    signoff = doc.add_paragraph()
    signoff.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signoff.paragraph_format.space_before = Pt(16)
    run = signoff.add_run(f"{AUTHOR}  |  {STUDENT_ID}")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
